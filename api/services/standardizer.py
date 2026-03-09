import os
import logging
from datetime import datetime
from django.conf import settings
from docx import Document

logger = logging.getLogger("api")


class ResumeStandardizer:


    def generate_docx(self, data: dict, output_filename: str = "") -> str:
    
        try:
            doc = Document()

            # Set document styles
            self._setup_styles(doc)

            # 1. Header
            self._add_header(doc, data)

            # 2. Professional Summary
            self._add_professional_summary(doc, data)

            # 3. Key Skills
            self._add_skills(doc, data)

            # 4. Work Experience
            self._add_work_experience(doc, data)

            # 5. Education
            self._add_education(doc, data)

            # 6. Certifications (optional)
            self._add_certifications(doc, data)

            # 7. Projects (optional)
            self._add_projects(doc, data)

            # Save document
            if not output_filename:
                name = data.get("candidate_full_name", "resume").replace(" ", "_")
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_filename = f"Standardized_Resume_{name}_{timestamp}.docx"

            output_path = os.path.join(settings.STANDARDIZED_OUTPUT_DIR, output_filename)
            doc.save(output_path)

            logger.info(f"Standardized resume generated: {output_path}")
            return output_path

        except ImportError:
            logger.error("python-docx is required for DOCX generation")
            raise
        except Exception as e:
            logger.error(f"Resume generation failed: {e}")
            raise

    def _setup_styles(self, doc):
        """Set up document styles for consistent formatting."""
        from docx.shared import Pt, RGBColor

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Reduce paragraph spacing
        paragraph_format = style.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(4)

    def _add_header(self, doc, data: dict):
        """Add resume header with name, contact info, and links."""
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Name
        name = data.get("candidate_full_name", "")
        if name:
            heading = doc.add_paragraph()
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = heading.add_run(name.upper())
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # Contact line
        contact_parts = []

        phones = data.get("phones", [])
        if phones:
            contact_parts.append(phones[0])

        emails = data.get("email_ids", [])
        if emails:
            contact_parts.append(emails[0])

        location = data.get("current_location", "")
        if location:
            contact_parts.append(location)

        linkedin = data.get("linkedin_url", "")
        if linkedin:
            contact_parts.append(linkedin)

        github = data.get("github_url", "")
        if github:
            contact_parts.append(github)

        if contact_parts:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = contact_para.add_run(" | ".join(contact_parts))
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Divider line
        divider = doc.add_paragraph()
        divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = divider.add_run("─" * 60)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    def _add_section_heading(self, doc, title: str):
        """Add a styled section heading."""
        from docx.shared import Pt, RGBColor

        heading = doc.add_paragraph()
        run = heading.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # Bottom border effect
        border_para = doc.add_paragraph()
        run = border_para.add_run("─" * 60)
        run.font.size = Pt(6)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    def _add_professional_summary(self, doc, data: dict):
        """Add professional summary section."""
        from docx.shared import Pt

        summary = data.get("professional_summary", "")
        if not summary:
            return

        self._add_section_heading(doc, "Professional Summary")
        para = doc.add_paragraph()
        run = para.add_run(summary)
        run.font.size = Pt(11)

    def _add_skills(self, doc, data: dict):
        """Add key skills section, grouped if possible."""
        from docx.shared import Pt

        skills = data.get("key_skills", [])
        if not skills:
            return

        self._add_section_heading(doc, "Key Skills")

        # Group skills by category
        skill_groups = self._categorize_skills(skills)

        for category, category_skills in skill_groups.items():
            if category != "Other":
                para = doc.add_paragraph()
                run = para.add_run(f"{category}: ")
                run.bold = True
                run.font.size = Pt(10)
                run = para.add_run(", ".join(category_skills))
                run.font.size = Pt(10)
            else:
                for skill in category_skills:
                    para = doc.add_paragraph(style="List Bullet")
                    run = para.add_run(skill)
                    run.font.size = Pt(10)

    def _categorize_skills(self, skills: list) -> dict:
        """Group skills into categories."""
        categories = {
            "Power Platform": [],
            "Programming": [],
            "Data & Analytics": [],
            "Cloud & DevOps": [],
            "Tools & Technologies": [],
            "Other": [],
        }

        power_platform_keywords = [
            "power apps", "power automate", "power bi", "power virtual",
            "copilot", "dataverse", "dynamics", "sharepoint",
        ]
        programming_keywords = [
            "python", "java", "javascript", "typescript", "c#", "c++",
            "react", "angular", "vue", "node", "django", "flask", ".net",
            "html", "css", "sql", "r ", "go", "rust", "swift", "kotlin",
        ]
        data_keywords = [
            "sql", "pandas", "numpy", "tableau", "excel", "data",
            "analytics", "machine learning", "ai", "deep learning",
            "tensorflow", "pytorch", "spark", "hadoop", "etl",
        ]
        cloud_keywords = [
            "azure", "aws", "gcp", "docker", "kubernetes", "ci/cd",
            "jenkins", "terraform", "devops", "git", "linux",
        ]

        for skill in skills:
            skill_lower = skill.lower().strip()
            categorized = False

            for keyword in power_platform_keywords:
                if keyword in skill_lower:
                    categories["Power Platform"].append(skill)
                    categorized = True
                    break
            if categorized:
                continue

            for keyword in programming_keywords:
                if keyword in skill_lower:
                    categories["Programming"].append(skill)
                    categorized = True
                    break
            if categorized:
                continue

            for keyword in data_keywords:
                if keyword in skill_lower:
                    categories["Data & Analytics"].append(skill)
                    categorized = True
                    break
            if categorized:
                continue

            for keyword in cloud_keywords:
                if keyword in skill_lower:
                    categories["Cloud & DevOps"].append(skill)
                    categorized = True
                    break
            if categorized:
                continue

            categories["Other"].append(skill)

        # Remove empty categories
        return {k: v for k, v in categories.items() if v}

    def _add_work_experience(self, doc, data: dict):
        """Add work experience section in reverse chronological order."""
        from docx.shared import Pt, RGBColor

        experiences = data.get("work_experience", [])
        if not experiences:
            return

        self._add_section_heading(doc, "Work Experience")

        for exp in experiences:
            # Company and Role line
            company = exp.get("company", "")
            role = exp.get("role", "")
            start = exp.get("start_date", "")
            end = exp.get("end_date", "Present")

            header_para = doc.add_paragraph()
            if company:
                run = header_para.add_run(company)
                run.bold = True
                run.font.size = Pt(11)
            if role:
                if company:
                    run = header_para.add_run(" — ")
                    run.font.size = Pt(11)
                run = header_para.add_run(role)
                run.italic = True
                run.font.size = Pt(11)

            # Dates
            if start or end:
                date_para = doc.add_paragraph()
                date_text = f"{start} – {end}"
                duration = exp.get("duration", "")
                if duration:
                    date_text += f" ({duration})"
                run = date_para.add_run(date_text)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

            # Responsibilities
            responsibilities = exp.get("responsibilities", [])
            # Fallback: legacy summary string
            if not responsibilities:
                summary = exp.get("summary", "")
                if summary:
                    responsibilities = [
                        b.strip().lstrip("•-– ")
                        for b in (summary.split("\n") if "\n" in summary else [summary])
                        if b.strip()
                    ]

            if responsibilities:
                for bullet in responsibilities:
                    if bullet.strip():
                        para = doc.add_paragraph(style="List Bullet")
                        run = para.add_run(bullet.strip().lstrip("•-– "))
                        run.font.size = Pt(10)

            # Achievements (shown after responsibilities, distinguished label)
            achievements = exp.get("achievements", [])
            if achievements:
                ach_label = doc.add_paragraph()
                run = ach_label.add_run("Key Achievements:")
                run.bold = True
                run.font.size = Pt(10)
                for ach in achievements:
                    if ach.strip():
                        para = doc.add_paragraph(style="List Bullet")
                        run = para.add_run(ach.strip().lstrip("•-– "))
                        run.font.size = Pt(10)

            # Technologies used
            technologies = exp.get("technologies", [])
            if technologies:
                tech_para = doc.add_paragraph()
                run = tech_para.add_run("Technologies: ")
                run.bold = True
                run.italic = True
                run.font.size = Pt(9)
                run = tech_para.add_run(", ".join(technologies))
                run.italic = True
                run.font.size = Pt(9)

            # Spacing between entries
            doc.add_paragraph()

    def _add_education(self, doc, data: dict):
        """Add education section."""
        from docx.shared import Pt

        education = data.get("education", [])
        if not education:
            return

        self._add_section_heading(doc, "Education")

        for edu in education:
            degree = edu.get("degree", "")
            # Support both extractor key (field_of_study) and legacy key (specialization)
            specialization = edu.get("field_of_study") or edu.get("specialization", "")
            institution = edu.get("institution", "")
            # Support both extractor key (end_year) and legacy key (year)
            year = edu.get("end_year") or edu.get("year", "")
            start_year = edu.get("start_year", "")
            grade = edu.get("grade", "")

            para = doc.add_paragraph()

            degree_text = degree
            if specialization:
                degree_text += f" in {specialization}"

            if degree_text:
                run = para.add_run(degree_text)
                run.bold = True
                run.font.size = Pt(11)

            if institution:
                if degree_text:
                    run = para.add_run(f"\n{institution}")
                else:
                    run = para.add_run(institution)
                run.font.size = Pt(10)

            year_range = ""
            if start_year and year:
                year_range = f"{start_year} – {year}"
            elif year:
                year_range = year

            if year_range:
                run = para.add_run(f" | {year_range}")
                run.font.size = Pt(10)

            if grade:
                run = para.add_run(f" | {grade}")
                run.font.size = Pt(10)

    def _add_certifications(self, doc, data: dict):
        """Add certifications section (optional)."""
        from docx.shared import Pt

        certifications = data.get("certifications", [])
        if not certifications:
            return

        self._add_section_heading(doc, "Certifications")

        for cert in certifications:
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(cert)
            run.font.size = Pt(10)

    def _add_projects(self, doc, data: dict):
        """Add projects section (optional, max 3)."""
        from docx.shared import Pt

        projects = data.get("projects", [])
        if not projects:
            return

        self._add_section_heading(doc, "Projects")

        for project in projects[:3]:  # Max 3 projects
            name = project.get("name", "")
            description = project.get("description", "")
            technologies = project.get("technologies", [])

            if name:
                para = doc.add_paragraph()
                run = para.add_run(name)
                run.bold = True
                run.font.size = Pt(11)

            if description:
                para = doc.add_paragraph()
                run = para.add_run(description)
                run.font.size = Pt(10)

            if technologies:
                para = doc.add_paragraph()
                run = para.add_run(f"Technologies: {', '.join(technologies)}")
                run.italic = True
                run.font.size = Pt(9)

    def generate_text_output(self, data: dict) -> str:
        """
        Generate a plain text version of the standardized resume.
        Used as fallback or for chat display.
        """
        lines = []

        # Header
        name = data.get("candidate_full_name", "")
        if name:
            lines.append(name.upper())
            lines.append("=" * len(name))

        contact = []
        if data.get("phones"):
            contact.append(data["phones"][0])
        if data.get("email_ids"):
            contact.append(data["email_ids"][0])
        if data.get("current_location"):
            contact.append(data["current_location"])
        if data.get("linkedin_url"):
            contact.append(data["linkedin_url"])
        if contact:
            lines.append(" | ".join(contact))

        lines.append("")

        # Professional Summary
        if data.get("professional_summary"):
            lines.append("PROFESSIONAL SUMMARY")
            lines.append("-" * 20)
            lines.append(data["professional_summary"])
            lines.append("")

        # Skills
        if data.get("key_skills"):
            lines.append("KEY SKILLS")
            lines.append("-" * 10)
            lines.append(", ".join(data["key_skills"]))
            lines.append("")

        # Experience
        if data.get("work_experience"):
            lines.append("WORK EXPERIENCE")
            lines.append("-" * 15)
            for exp in data["work_experience"]:
                company = exp.get("company", "")
                role = exp.get("role", "")
                start = exp.get("start_date", "")
                end = exp.get("end_date", "Present")
                duration = exp.get("duration", "")
                date_str = f"{start} – {end}" + (f" ({duration})" if duration else "")
                lines.append(f"{company} — {role} | {date_str}")
                responsibilities = exp.get("responsibilities", [])
                if not responsibilities and exp.get("summary"):
                    summary = exp["summary"]
                    responsibilities = [
                        b.strip().lstrip("•-– ")
                        for b in (summary.split("\n") if "\n" in summary else [summary])
                        if b.strip()
                    ]
                for bullet in responsibilities:
                    if bullet.strip():
                        lines.append(f"  • {bullet.strip().lstrip('•-– ')}")
                for ach in exp.get("achievements", []):
                    if ach.strip():
                        lines.append(f"  ★ {ach.strip().lstrip('•-– ')}")
                techs = exp.get("technologies", [])
                if techs:
                    lines.append(f"  Technologies: {', '.join(techs)}")
                lines.append("")

        # Education
        if data.get("education"):
            lines.append("EDUCATION")
            lines.append("-" * 9)
            for edu in data["education"]:
                degree = edu.get("degree", "")
                specialization = edu.get("field_of_study") or edu.get("specialization", "")
                institution = edu.get("institution", "")
                year = edu.get("end_year") or edu.get("year", "")
                start_year = edu.get("start_year", "")
                grade = edu.get("grade", "")
                degree_text = f"{degree} in {specialization}" if specialization else degree
                year_range = f"{start_year} – {year}" if start_year and year else year
                edu_line = " | ".join(p for p in [degree_text, institution, year_range, grade] if p)
                lines.append(edu_line)
            lines.append("")

        # Certifications
        if data.get("certifications"):
            lines.append("CERTIFICATIONS")
            lines.append("-" * 14)
            for cert in data["certifications"]:
                lines.append(f"  • {cert}")
            lines.append("")

        return "\n".join(lines)
